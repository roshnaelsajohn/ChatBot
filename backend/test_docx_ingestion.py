
import os
import sys
from docx import Document

# Add current directory to path
sys.path.append('/app')

from rag_service import get_rag_service

def test_docx():
    print("🚀 Starting DOCX Ingestion Test...")
    
    # 1. Create a dummy DOCX
    doc = Document()
    doc.add_paragraph("This is a test document in DOCX format.")
    doc.add_paragraph("It contains specific information about Apple Pie recipes.")
    doc.add_paragraph("Apple Pie requires apples, sugar, and crust.")
    
    test_filename = "test_recipe.docx"
    test_path = f"/tmp/{test_filename}"
    doc.save(test_path)
    
    print(f"📄 Created test DOCX: {test_path}")

    # 2. Store Document
    rag_service = get_rag_service()
    print("📥 Storing document...")
    try:
        result = rag_service.process_and_store_document(test_path, test_filename, "docx")
        print(f"Store Result: {result}")
    except Exception as e:
        print(f"❌ Exception storing document: {e}")

    # 3. Verify
    stats = rag_service.get_collection_stats()
    print(f"📊 Current Stats: {stats}")

if __name__ == "__main__":
    test_docx()
